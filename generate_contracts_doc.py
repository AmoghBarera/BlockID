import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('BlockID: Smart Contracts Deep Dive & Interview QA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    
    doc.add_paragraph('This document is a comprehensive guide to everything inside the "contracts" folder. Use this to prepare for any technical interview questions focused on Solidity, Ethereum, and Smart Contract Architecture.', style='Intense Quote')
    
    # 1. Structure of the Contracts Folder
    doc.add_heading('1. Structure of the "contracts" Folder', level=1)
    doc.add_paragraph('The contracts directory is a Hardhat project. Hardhat is an Ethereum development environment used by professionals to compile, deploy, test, and debug smart contracts.')
    p_structure = doc.add_paragraph()
    p_structure.add_run('contracts/:').bold = True
    p_structure.add_run(' Contains the Solidity source code (`IdentityRegistry.sol`).\n')
    p_structure.add_run('scripts/:').bold = True
    p_structure.add_run(' Contains deployment logic written in TypeScript (`deploy.ts`). This script uses the Ethers.js library to push the compiled bytecode onto the blockchain.\n')
    p_structure.add_run('hardhat.config.ts:').bold = True
    p_structure.add_run(' The master configuration file. It sets the Solidity compiler version (0.8.24), configures the local node (`localhost:8545`), and sets up testnets like Polygon Amoy.')
    
    # 2. Deep Dive: IdentityRegistry.sol
    doc.add_heading('2. Deep Dive: How IdentityRegistry.sol Works', level=1)
    doc.add_paragraph('This is the core ledger of the BlockID application. It is written in Solidity.')
    
    doc.add_heading('Role-Based Access Control (RBAC)', level=2)
    doc.add_paragraph('The contract inherits from OpenZeppelin\'s `AccessControl`. OpenZeppelin is a standard library for secure contract development. This allows us to define specific roles instead of giving everyone equal power:')
    doc.add_paragraph('- DEFAULT_ADMIN_ROLE: The master admin (who deployed the contract).')
    doc.add_paragraph('- VERIFIER_ROLE: Authorized institutions (like Banks/Colleges) who are allowed to mark a user as "Verified".')
    
    doc.add_heading('Core Data Structures', level=2)
    doc.add_paragraph('The contract uses two main data structures (structs):')
    doc.add_paragraph('1. IdentityRecord: Stores a user\'s Wallet Address, DID (Decentralized Identifier), IPFS Metadata CID (where the real docs live), Proof Hash, and their current Status (Registered, Verified, etc.).')
    doc.add_paragraph('2. VerificationEntry: A historic log of who verified the user, keeping an audit trail on the blockchain.')
    doc.add_paragraph('It maps this data so it can be quickly looked up: `mapping(string => IdentityRecord) private identities;`')
    
    doc.add_heading('Key Functions', level=2)
    doc.add_paragraph('- registerIdentity(): Allows any new wallet to register a DID. Ensures a wallet cannot register twice.')
    doc.add_paragraph('- requestVerification(): Allows a user to change their status to "PendingVerification".')
    doc.add_paragraph('- verifyIdentity(): ONLY callable by a wallet with the VERIFIER_ROLE. It changes the status to Verified or Rejected and adds an entry to the audit log.')
    doc.add_paragraph('- updateIdentity(): Allows users to update their IPFS link if their real-world documents change.')
    doc.add_paragraph('- grantVerifierRole() / revokeVerifierRole(): ONLY callable by ADMIN. Grants institutional wallets the power to verify users.')

    # 3. Smart Contract Interview Questions
    doc.add_heading('3. Smart Contract Interview / Viva Questions', level=1)
    
    doc.add_paragraph('Q1. What is Hardhat and why did you use it instead of Remix?').bold = True
    doc.add_paragraph('Answer: "Remix is an excellent web-based IDE for simple prototyping, but Hardhat is an industry-standard local development environment. We used Hardhat because it provides a local lightweight Ethereum node for testing, allows us to write deployment and test scripts in TypeScript, and integrates seamlessly with our total frontend/backend repository."')

    doc.add_paragraph('Q2. What are "Gas Fees" and how did you optimize for them in your contract?').bold = True
    doc.add_paragraph('Answer: "Gas is the execution fee paid to miners/validators to process a transaction on the EVM (Ethereum Virtual Machine). Since writing data to the blockchain is very expensive, we optimized gas by NOT storing strings of raw data (like names/addresses) on the blockchain. Instead, we only store fixed-size cryptographic Hashes and IPFS CIDs. We also enabled the Solidity Optimizer in hardhat.config.ts."')

    doc.add_paragraph('Q3. Why did you use OpenZeppelin\'s AccessControl?').bold = True
    doc.add_paragraph('Answer: "Security. We needed to guarantee that normal users cannot just verify themselves. By inheriting AccessControl from OpenZeppelin, we get battle-tested, secure logic that assigns a VERIFIER_ROLE to specific wallets. If a hacker tries to call verifyIdentity(), the transaction will violently fail and revert because they lack the cryptographic role hash."')

    doc.add_paragraph('Q4. What is the difference between mapping and an array in Solidity, and why use mapping for identities?').bold = True
    doc.add_paragraph('Answer: "An array is a list of items where iterating to find a specific item costs a massive amount of gas as the list grows (O(N) time). A mapping is like a hash table or dictionary. It allows us to instantly look up an identity in O(1) time utilizing the DID as the key (`mapping(string => IdentityRecord)`). This keeps transaction gas costs uniformly low regardless of how many users register."')

    doc.add_paragraph('Q5. What happens if someone uploads fake documents during `registerIdentity()`?').bold = True
    doc.add_paragraph('Answer: "The smart contract cannot read real-world documents; it is blind to the outside world. If they upload fake documents, the smart contract blindly registers the hash. However, the system is secure because that user\'s status remains just \'Registered\'. When an authorized VERIFIER checks the fake documents off-chain during `verifyIdentity()`, they will see they are fake and call the contract to change the status to \'Rejected\'."')

    doc.add_paragraph('Q6. Why are the identities mapped using `private` visibility? `mapping(string => IdentityRecord) private identities;`').bold = True
    doc.add_paragraph('Answer: "Nothing on the blockchain is truly private; all data can be read. However, marking the mapping as `private` prevents Solidity from automatically generating a simple getter function. We instead wrote a custom `getIdentity(did)` function so we can provide custom reverting error messages (like \'IdentityNotFound\') if someone searches for a DID that doesn\'t exist, rather than returning empty default zero values."')

    doc.add_paragraph('Q7. How do you deploy this contract to a real network like Polygon?').bold = True
    doc.add_paragraph('Answer: "Inside our `hardhat.config.ts`, we add a network configuration containing an RPC URL (like Infura/Alchemy) and the private key of an Administration wallet holding real testnet MATIC tokens. Then, running `npx hardhat run scripts/deploy.ts --network amoy` securely signs the bytecode and pushes it to the real network."')

    doc.save('BlockID_Contracts_DeepDive.docx')
    print("Report generated successfully as 'BlockID_Contracts_DeepDive.docx'")

if __name__ == "__main__":
    main()
